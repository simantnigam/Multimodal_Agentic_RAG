from __future__ import annotations

import base64
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from core.constants import (
    DOC_TYPE_PDF,
    DOC_TYPE_HTML,
    DOC_TYPE_DOCX,
    DOC_TYPE_IMAGE,
    MODALITY_IMAGE,
    MODALITY_TABLE,
    VISION_PROVIDER_ANTHROPIC,
    VISION_PROVIDER_OPENAI,
    VISION_PROVIDER_NONE,
)
from core.logging import get_logger

logger = get_logger(__name__)

# Lazy module-level vision clients — created once, reused across all images in a parse run.
_anthropic_client = None
_openai_client    = None

_PDF_HEADING_FONT_SIZE = 14

_CAPTION_PROMPT = (
    "Describe this image concisely for a document retrieval system. "
    "Focus on the key information, data, or concepts shown. "
    "If it is a chart or graph, describe the data and trends. "
    "If it is a diagram, describe the structure and relationships. "
    "Respond in 2-4 sentences."
)

_MIME = {
    "jpg": "image/jpeg", "jpeg": "image/jpeg",
    "png": "image/png",  "gif": "image/gif",
    "webp": "image/webp",
}


# ---------------------------------------------------------------------------
# Data contracts
# ---------------------------------------------------------------------------

@dataclass
class ParsedImage:
    caption:  str
    data:     bytes
    fmt:      str
    page:     int = 0
    modality: str = MODALITY_IMAGE


@dataclass
class ParsedTable:
    content:  str
    page:     int = 0
    modality: str = MODALITY_TABLE


@dataclass
class ParsedSection:
    title:  str
    level:  int
    blocks: list[str] = field(default_factory=list)


@dataclass
class ParsedDocument:
    title:    str
    source:   str
    doc_type: str
    sections: list[ParsedSection] = field(default_factory=list)
    images:   list[ParsedImage]   = field(default_factory=list)
    tables:   list[ParsedTable]   = field(default_factory=list)
    metadata: dict[str, Any]      = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Vision captioning
# ---------------------------------------------------------------------------

def caption_image(data: bytes, fmt: str) -> str:
    """
    Generate a semantic caption for an image using a configured vision LLM.

    Provider selected via VISION_PROVIDER in .env:
        anthropic → Claude (claude-sonnet-4-6 or configured VISION_MODEL)
        openai    → GPT-4o (or configured VISION_MODEL)
        none      → returns empty string (no captioning)

    Falls back to empty string on any API failure — ingestion is not blocked.
    """
    from core.config import get_settings
    settings = get_settings()
    provider = settings.vision_provider.lower()

    if provider == VISION_PROVIDER_NONE or not provider:
        return ""

    try:
        if provider == VISION_PROVIDER_ANTHROPIC:
            return _caption_anthropic(data, fmt, settings)
        if provider == VISION_PROVIDER_OPENAI:
            return _caption_openai(data, fmt, settings)
        logger.warning("unknown_vision_provider", provider=provider)
        return ""
    except Exception as exc:
        logger.warning("vision_captioning_failed", provider=provider, error=str(exc))
        return ""


def _caption_anthropic(data: bytes, fmt: str, settings) -> str:
    global _anthropic_client
    import anthropic

    if _anthropic_client is None:
        _anthropic_client = anthropic.Anthropic(api_key=settings.anthropic_api_key)

    mime_type = _MIME.get(fmt.lower(), "image/jpeg")
    b64       = base64.standard_b64encode(data).decode("utf-8")

    response = _anthropic_client.messages.create(
        model=settings.vision_model,
        max_tokens=settings.vision_max_tokens,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image",
                    "source": {
                        "type":       "base64",
                        "media_type": mime_type,
                        "data":       b64,
                    },
                },
                {"type": "text", "text": _CAPTION_PROMPT},
            ],
        }],
    )
    return response.content[0].text.strip()


def _caption_openai(data: bytes, fmt: str, settings) -> str:
    global _openai_client
    import openai

    if _openai_client is None:
        _openai_client = openai.OpenAI(api_key=settings.openai_api_key)

    mime_type = _MIME.get(fmt.lower(), "image/jpeg")
    b64       = base64.standard_b64encode(data).decode("utf-8")
    data_url  = f"data:{mime_type};base64,{b64}"

    response = _openai_client.chat.completions.create(
        model=settings.vision_model,
        max_tokens=settings.vision_max_tokens,
        messages=[{
            "role": "user",
            "content": [
                {"type": "image_url", "image_url": {"url": data_url}},
                {"type": "text",      "text": _CAPTION_PROMPT},
            ],
        }],
    )
    return response.choices[0].message.content.strip()


# ---------------------------------------------------------------------------
# Public dispatcher
# ---------------------------------------------------------------------------

def parse_document(path: str | Path) -> ParsedDocument:
    """
    Parse a document from disk and return a structured ParsedDocument.
    Images are captioned via the configured vision LLM (VISION_PROVIDER).

    Dispatches by file extension:
        .pdf            → PyMuPDF
        .html / .htm    → stdlib HTMLParser
        .docx           → python-docx
        .jpg/jpeg/png   → vision caption only (no text blocks)
    """
    path = Path(path)
    ext  = path.suffix.lower()

    logger.info("parsing_document", path=str(path), ext=ext)

    if ext == ".pdf":
        return _parse_pdf(path)
    if ext in {".html", ".htm"}:
        return _parse_html(path)
    if ext == ".docx":
        return _parse_docx(path)
    if ext in {".jpg", ".jpeg", ".png", ".gif", ".webp"}:
        return _parse_image(path)

    raise ValueError(
        f"Unsupported file type '{ext}'. Supported: pdf, html, docx, jpg/png."
    )


# ---------------------------------------------------------------------------
# PDF parser
# ---------------------------------------------------------------------------

def _parse_pdf(path: Path) -> ParsedDocument:
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    title    = doc.metadata.get("title") or path.stem
    sections: list[ParsedSection] = []
    images:   list[ParsedImage]   = []
    tables:   list[ParsedTable]   = []

    current_section = ParsedSection(title=title, level=1)

    for page_num, page in enumerate(doc):
        # --- Text blocks ---
        for block in page.get_text("dict")["blocks"]:
            if "lines" not in block:
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    text      = span["text"].strip()
                    font_size = span.get("size", 0)
                    is_bold   = "Bold" in span.get("font", "")
                    if not text:
                        continue
                    if font_size >= _PDF_HEADING_FONT_SIZE or (is_bold and font_size >= 11):
                        if current_section.blocks:
                            sections.append(current_section)
                        level = 1 if font_size >= 18 else 2
                        current_section = ParsedSection(title=text, level=level)
                    else:
                        current_section.blocks.append(text)

        # --- Images (with vision captioning) ---
        for img in page.get_images(full=True):
            xref       = img[0]
            base_image = doc.extract_image(xref)
            img_data   = base_image["image"]
            img_fmt    = base_image["ext"]
            caption    = caption_image(img_data, img_fmt)
            if not caption:
                caption = f"Figure on page {page_num + 1}"
            images.append(ParsedImage(
                caption=caption,
                data=img_data,
                fmt=img_fmt,
                page=page_num + 1,
            ))

        # --- Tables ---
        for tab in page.find_tables():
            df = tab.to_pandas()
            try:
                md = df.to_markdown(index=False)
            except ImportError:
                md = df.to_string(index=False)
            tables.append(ParsedTable(content=md, page=page_num + 1))

    if current_section.blocks:
        sections.append(current_section)

    logger.info(
        "pdf_parsed",
        path=str(path),
        pages=len(doc),
        sections=len(sections),
        images=len(images),
        tables=len(tables),
    )

    return ParsedDocument(
        title=title,
        source=str(path),
        doc_type=DOC_TYPE_PDF,
        sections=sections,
        images=images,
        tables=tables,
        metadata=dict(doc.metadata),
    )


# ---------------------------------------------------------------------------
# HTML parser
# ---------------------------------------------------------------------------

class _HTMLExtractor(HTMLParser):
    _HEADING_TAGS = {"h1", "h2", "h3", "h4"}
    _SKIP_TAGS    = {"script", "style", "head", "meta", "link", "nav", "footer"}

    def __init__(self) -> None:
        super().__init__()
        self.sections: list[ParsedSection] = []
        self._current_section = ParsedSection(title="Document", level=1)
        self._current_tag  = ""
        self._skip         = False
        self._buffer       = ""

    def handle_starttag(self, tag: str, attrs: list) -> None:
        self._current_tag = tag
        self._skip        = tag in self._SKIP_TAGS
        self._buffer      = ""

    def handle_endtag(self, tag: str) -> None:
        text = self._buffer.strip()
        if not text or self._skip:
            return
        if tag in self._HEADING_TAGS:
            if self._current_section.blocks:
                self.sections.append(self._current_section)
            level = int(tag[1])
            self._current_section = ParsedSection(title=text, level=level)
        elif tag in {"p", "li", "td", "th"}:
            if text:
                self._current_section.blocks.append(text)
        self._buffer = ""

    def handle_data(self, data: str) -> None:
        if not self._skip:
            self._buffer += data

    def finish(self) -> list[ParsedSection]:
        if self._current_section.blocks:
            self.sections.append(self._current_section)
        return self.sections


def _parse_html(path: Path) -> ParsedDocument:
    html      = path.read_text(encoding="utf-8", errors="ignore")
    extractor = _HTMLExtractor()
    extractor.feed(html)
    sections  = extractor.finish()
    title     = sections[0].title if sections else path.stem

    logger.info("html_parsed", path=str(path), sections=len(sections))

    return ParsedDocument(
        title=title,
        source=str(path),
        doc_type=DOC_TYPE_HTML,
        sections=sections,
    )


# ---------------------------------------------------------------------------
# DOCX parser
# ---------------------------------------------------------------------------

_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title"}


def _parse_docx(path: Path) -> ParsedDocument:
    from docx import Document as DocxDocument

    docx    = DocxDocument(str(path))
    sections: list[ParsedSection] = []
    tables:   list[ParsedTable]   = []
    title   = path.stem
    current = ParsedSection(title=title, level=1)

    for para in docx.paragraphs:
        text  = para.text.strip()
        style = para.style.name if para.style else ""
        if not text:
            continue
        if style in _HEADING_STYLES or style.startswith("Heading"):
            if current.blocks:
                sections.append(current)
            level   = 1 if "1" in style or style == "Title" else 2
            current = ParsedSection(title=text, level=level)
            if style == "Title":
                title = text
        else:
            current.blocks.append(text)

    if current.blocks:
        sections.append(current)

    for table in docx.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            header = " | ".join(rows[0])
            sep    = " | ".join(["---"] * len(rows[0]))
            body   = "\n".join(" | ".join(r) for r in rows[1:])
            tables.append(ParsedTable(content=f"{header}\n{sep}\n{body}"))

    logger.info("docx_parsed", path=str(path), sections=len(sections), tables=len(tables))

    return ParsedDocument(
        title=title,
        source=str(path),
        doc_type=DOC_TYPE_DOCX,
        sections=sections,
        tables=tables,
    )


# ---------------------------------------------------------------------------
# Image parser — single image file, captioned via vision LLM
# ---------------------------------------------------------------------------

def _parse_image(path: Path) -> ParsedDocument:
    data    = path.read_bytes()
    fmt     = path.suffix.lstrip(".").lower()
    caption = caption_image(data, fmt)
    if not caption:
        caption = path.stem

    images   = [ParsedImage(caption=caption, data=data, fmt=fmt)]
    sections = [ParsedSection(title=path.stem, level=1, blocks=[])]

    logger.info("image_parsed", path=str(path), size_bytes=len(data), captioned=bool(caption))

    return ParsedDocument(
        title=path.stem,
        source=str(path),
        doc_type=DOC_TYPE_IMAGE,
        sections=sections,
        images=images,
    )
